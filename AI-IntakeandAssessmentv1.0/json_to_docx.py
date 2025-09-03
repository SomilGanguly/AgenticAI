import argparse
import json
import os
import sys
from typing import Any, Dict, List, Optional


def _try_import_python_docx():
    try:
        from docx import Document  # type: ignore
        return Document
    except Exception:
        print("Missing dependency: python-docx. Install it first (e.g., pip install python-docx)", file=sys.stderr)
        raise


def _safe_text(v: Any) -> str:
    return "" if v is None else str(v)


def _clamp_heading_level(level: Optional[int]) -> int:
    try:
        lv = int(level) if level is not None else 1
    except Exception:
        lv = 1
    # python-docx supports 0..9 (0 is Title); we map 1..9 here
    return max(1, min(9, lv))


def json_to_docx(data: Dict[str, Any], out_path: str) -> str:
    Document = _try_import_python_docx()
    doc = Document()

    blocks = data.get("blocks") or []
    if not isinstance(blocks, list):
        blocks = []

    for blk in blocks:
        if not isinstance(blk, dict):
            continue
        btype = (blk.get("type") or "paragraph").lower()
        if btype == "heading":
            text = _safe_text(blk.get("text"))
            level = _clamp_heading_level(blk.get("level"))
            if text:
                # python-docx: level 0..9 (0=Title). We use 1..9.
                doc.add_heading(text, level=level)
        elif btype == "table":
            rows = blk.get("rows") or []
            if not isinstance(rows, list) or len(rows) == 0:
                continue
            # determine column count (max row length)
            col_count = max((len(r) for r in rows if isinstance(r, list)), default=0)
            if col_count == 0:
                continue
            table = doc.add_table(rows=len(rows), cols=col_count)
            for r_idx, r in enumerate(rows):
                if not isinstance(r, list):
                    r = []
                # pad row to col_count
                padded = list(r[:col_count]) + ["" for _ in range(max(0, col_count - len(r)))]
                for c_idx in range(col_count):
                    table.cell(r_idx, c_idx).text = _safe_text(padded[c_idx])
        else:  # paragraph (default)
            text = _safe_text(blk.get("text"))
            if text:
                doc.add_paragraph(text)

    # ensure output directory
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def convert_file(json_path: str, out_dir: str) -> str:
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    base = os.path.splitext(os.path.basename(json_path))[0]
    out_path = os.path.join(out_dir, base + ".docx")
    return json_to_docx(data, out_path)


def convert_folder(input_dir: str, out_dir: str) -> List[str]:
    outputs: List[str] = []
    for name in os.listdir(input_dir):
        if not name.lower().endswith(".json"):
            continue
        in_path = os.path.join(input_dir, name)
        if not os.path.isfile(in_path):
            continue
        outputs.append(convert_file(in_path, out_dir))
    return outputs


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert JSON (from docx_to_json.py) back to Word .docx")
    g = parser.add_mutually_exclusive_group(required=True)
    g.add_argument("--input-file", help="Path to a single JSON file to convert")
    g.add_argument("--input-dir", help="Path to a directory containing JSON files to convert")
    parser.add_argument("--output-dir", default="out_docx", help="Output directory for .docx files (default: out_docx)")
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)

    try:
        if args.input_file:
            if not os.path.isfile(args.input_file):
                print(f"File not found: {args.input_file}", file=sys.stderr)
                sys.exit(1)
            out_path = convert_file(args.input_file, args.output_dir)
            print(out_path)
        else:
            if not os.path.isdir(args.input_dir):
                print(f"Directory not found: {args.input_dir}", file=sys.stderr)
                sys.exit(1)
            outputs = convert_folder(args.input_dir, args.output_dir)
            for p in outputs:
                print(p)
    except Exception as ex:
        print(f"Conversion failed: {ex}", file=sys.stderr)
        sys.exit(2)


if __name__ == "__main__":
    main()
