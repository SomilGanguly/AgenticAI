import argparse
import json
import os
import re
import sys
from typing import Any, Dict, List, Optional


def _try_import_python_docx():
    try:
        from docx import Document  # type: ignore
        return Document
    except Exception as ex:
        print(
            "Missing dependency: python-docx. Install it first (e.g., pip install python-docx)",
            file=sys.stderr,
        )
        raise


def normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def heading_level_from_style(style_name: Optional[str]) -> Optional[int]:
    if not style_name:
        return None
    m = re.match(r"Heading\s*(\d+)", str(style_name), flags=re.IGNORECASE)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    return None


def extract_blocks_from_docx(path: str, include_empty: bool = False) -> Dict[str, Any]:
    Document = _try_import_python_docx()
    doc = Document(path)

    blocks: List[Dict[str, Any]] = []

    # paragraphs (includes headings)
    for p in doc.paragraphs:
        text = normalize_ws(p.text)
        level = heading_level_from_style(getattr(getattr(p, "style", None), "name", None))
        if level is not None:
            if text or include_empty:
                blocks.append({"type": "heading", "level": level, "text": text})
        else:
            if text or include_empty:
                blocks.append({"type": "paragraph", "text": text})

    # tables
    for t in doc.tables:
        rows: List[List[str]] = []
        for r in t.rows:
            row_vals: List[str] = []
            for c in r.cells:
                row_vals.append(normalize_ws(c.text))
            rows.append(row_vals)
        blocks.append({"type": "table", "rows": rows})

    return {"file": os.path.basename(path), "blocks": blocks}


def convert_folder(input_dir: str, output_dir: str, include_empty: bool = False) -> List[str]:
    if not os.path.isdir(input_dir):
        raise FileNotFoundError(f"Input directory not found: {input_dir}")
    os.makedirs(output_dir, exist_ok=True)

    outputs: List[str] = []
    for name in os.listdir(input_dir):
        if not name.lower().endswith(".docx"):
            continue
        in_path = os.path.join(input_dir, name)
        if not os.path.isfile(in_path):
            continue
        data = extract_blocks_from_docx(in_path, include_empty=include_empty)
        out_name = os.path.splitext(name)[0] + ".json"
        out_path = os.path.join(output_dir, out_name)
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        outputs.append(out_path)
    return outputs


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Word .docx files in a folder to JSON.")
    g = parser.add_mutually_exclusive_group(required=True)
    g.add_argument("--input-dir", help="Folder containing .docx files to convert")
    g.add_argument("--file", help="Single .docx file to convert")
    parser.add_argument("--output-dir", default="out_json", help="Folder to write JSON outputs (default: out_json)")
    parser.add_argument("--include-empty", action="store_true", help="Include empty paragraphs/headings")
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)

    if args.file:
        if not os.path.isfile(args.file):
            print(f"File not found: {args.file}", file=sys.stderr)
            sys.exit(1)
        data = extract_blocks_from_docx(args.file, include_empty=args.include_empty)
        base = os.path.splitext(os.path.basename(args.file))[0]
        out_path = os.path.join(args.output_dir, base + ".json")
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(out_path)
        return

    # folder mode
    outputs = convert_folder(args.input_dir, args.output_dir, include_empty=args.include_empty)
    for p in outputs:
        print(p)


if __name__ == "__main__":
    main()
